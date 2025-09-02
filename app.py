import os
import re
import datetime
import pdfplumber
import streamlit as st
from docx import Document
from openpyxl import load_workbook
from collections import defaultdict
import tempfile
import traceback
import shutil
from pathlib import Path

# 设置页面标题和布局
st.set_page_config(page_title="商标案件请款系统", layout="wide")
st.title("商标案件请款系统")
st.caption("案件类目前仅支持驳回复审、异议申请、无效申请和撤三申请")

# 初始化session状态
if 'processing_stage' not in st.session_state:
    st.session_state.processing_stage = 0  # 0: 未开始, 1: 提取完成, 2: 生成完成
if 'case_type' not in st.session_state:
    st.session_state.case_type = "新申请商标"  # 默认选择
if 'extracted_data' not in st.session_state:
    st.session_state.extracted_data = None
if 'agent_fees' not in st.session_state:
    st.session_state.agent_fees = {}
if 'generated_files' not in st.session_state:
    st.session_state.generated_files = []
if 'temp_dir' not in st.session_state:
    st.session_state.temp_dir = ""

# 官费标准
OFFICIAL_FEES = {
    "驳回复审": 675,
    "商标异议": 450,
    "撤三申请": 450,
    "无效宣告": 750,
    "新申请商标": 270,  # 新申请商标的官费
}

# 金额转大写函数
CN_NUM = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']
CN_UNIT = ['', '拾', '佰', '仟', '万', '拾', '佰', '仟', '亿']

def number_to_upper(amount):
    s = str(int(amount))
    result = []
    for i, ch in enumerate(s[::-1]):
        if int(ch) != 0:
            result.append(f"{CN_NUM[int(ch)]}{CN_UNIT[i]}")
    return ''.join(reversed(result)) + "元整"

# ============================= 新申请商标处理函数 =============================
def extract_pdf_data(pdf_path):
    """从新申请PDF提取数据"""
    applicant = "N/A"
    unified_credit_code = "N/A"
    final_date = "N/A"
    trademarks_with_categories = []
    pending_categories = []
    
    with pdfplumber.open(pdf_path) as pdf:
        all_texts = [page.extract_text().replace("　", " ").replace("\xa0", " ").strip() 
                     if page.extract_text() else "" for page in pdf.pages]
        all_text_combined = "\n---PAGE_BREAK---\n".join(all_texts)
        pages = all_text_combined.split("\n---PAGE_BREAK---\n")
        
        for page_num, page_text in enumerate(pages):
            # 第一页：提取申请人和统一社会信用代码
            if page_num == 0:
                applicant_match = re.search(r"申请人名称\(中文\)：\s*(.*?)\s*\(\s*英文\)", page_text)
                applicant = applicant_match.group(1).strip() if applicant_match else "N/A"
                
                # 使用统一的信用代码提取正则表达式
                unified_credit_code_match = re.search(r'(?:统一社会信用代码|信用代码)[：:]\s*([0-9A-Z]{18})', page_text, re.IGNORECASE)
                unified_credit_code = unified_credit_code_match.group(1).strip() if unified_credit_code_match else "N/A"
                
                # 尝试从第一页提取日期
                if final_date == "N/A":
                    date_match = re.search(r"(\d{4}年\s*\d{1,2}月\s*\d{1,2}日)", page_text)
                    final_date = date_match.group(1).replace(" ", "") if date_match else "N/A"
                continue
            
            # 后续页面：提取类别或商标名
            # 检查是否包含类别信息
            if re.search(r'类别：\d+', page_text):
                categories_found = re.findall(r'类别：(\d+)', page_text)
                pending_categories.extend(categories_found)
            
            # 检查是否包含委托书
            elif '商 标 代 理 委 托 书' in page_text:
                tm_name_match = re.search(r'商标代理委托书.*?代理\s+(.*?)商标\s*的\s*如下.*?事宜', 
                                         page_text, re.DOTALL)
                tm_name = tm_name_match.group(1).strip() if tm_name_match else ""
                
                if not tm_name:
                    fallback_match = re.search(r'代理\s+(.*?)\s*商标', page_text)
                    tm_name = fallback_match.group(1).strip() if fallback_match else ""
                
                if not tm_name:
                    st.warning(f"警告：在文件 {os.path.basename(pdf_path)} 的第 {page_num + 1} 页委托书中未找到商标名称。")
                
                # 提取委托书日期
                date_match = re.search(r"(\d{4}年\s*\d{1,2}月\s*\d{1,2}日)", page_text)
                if date_match:
                    final_date = date_match.group(1).replace(" ", "")
                
                # 关联类别与商标名
                if pending_categories:
                    for category in pending_categories:
                        trademarks_with_categories.append({
                            "商标名称": tm_name,
                            "类别": category
                        })
                    pending_categories.clear()
                else:
                    trademarks_with_categories.append({
                        "商标名称": tm_name,
                        "类别": "MANUAL_INPUT_REQUIRED"
                    })
                    st.warning(f"提示：文件 {os.path.basename(pdf_path)} 中的商标 '{tm_name}' 未找到自动关联的类别，需要手动输入。")
        
        # 检查是否还有未关联的类别
        if pending_categories:
            st.warning(f"警告：文件 {os.path.basename(pdf_path)} 处理完毕，但仍有未关联的类别 {pending_categories}。这些类别将被忽略。")
    
    return {
        "申请人": applicant,
        "统一社会信用代码": unified_credit_code,
        "日期": final_date,
        "商标列表": trademarks_with_categories,
        "事宜类型": "商标注册申请"
    }

# ============================= 案件类商标处理函数 =============================
def extract_case_info(text, filename):
    if any(kw in filename for kw in ['驳回', '复审']):
        return extract_review_case(text, filename)
    elif any(kw in filename for kw in ['撤三', '撤销连续']):
        return extract_non_use_case(text, filename)
    elif '异议' in filename:
        return extract_opposition_case(text, filename)
    elif any(kw in filename for kw in ['无效', '宣告']):
        return extract_invalid_case(text, filename)
    else:
        raise ValueError(f"无法识别案件类型: {filename}")

def extract_review_case(text, filename):
    case_type = "驳回复审"
    applicant = re.search(r'(?:申请人名称\$\$中文\$\$|申请人名称)：\s*([^\n]*?)(?=\s+(?:统一社会信用代码|地址))', 
                          text, re.DOTALL)
    applicant = applicant.group(1).strip() if applicant else "N/A"
    
    # 提取统一社会信用代码
    unified_credit_code_match = re.search(r'(?:统一社会信用代码|信用代码)[：:]\s*([0-9A-Z]{18})', text, re.IGNORECASE)
    unified_credit_code = unified_credit_code_match.group(1).strip() if unified_credit_code_match else "N/A"
    
    trademarks = []
    for m in re.finditer(r'申请商标：\s*(.*?)\s+类别：\s*(\d+).*?申请号/国际注册号：\s*([0-9A-Za-z]+)', 
                         text, re.DOTALL):
        trademarks.append({
            "商标名称": m.group(1).strip(), 
            "类别": int(m.group(2)), 
            "注册号": m.group(3)
        })
    
    return {
        "文件名": filename, 
        "案件类型": case_type, 
        "申请人": applicant,
        "统一社会信用代码": unified_credit_code,
        "商标列表": trademarks
    }

def extract_non_use_case(text, filename):
    case_type = "撤三申请"
    applicant = re.search(r'(?:申请人名称|申请人)：\s*([^\n]*?)(?=\s+(?:统一社会信用代码|地址))', 
                          text, re.DOTALL)
    applicant = applicant.group(1).strip() if applicant else "N/A"
    
    # 提取统一社会信用代码
    unified_credit_code_match = re.search(r'(?:统一社会信用代码|信用代码)[：:]\s*([0-9A-Z]{18})', text, re.IGNORECASE)
    unified_credit_code = unified_credit_code_match.group(1).strip() if unified_credit_code_match else "N/A"
    
    trademarks = []
    trademark_name_match = re.search(r'商标：\s*(.*?)\s*(?=\n|$)', text)
    category_match = re.search(r'类别：\s*(\d+)', text)
    registration_number_match = re.search(r'商标注册号：\s*([0-9A-Za-z]+)', text)
    
    if trademark_name_match and category_match and registration_number_match:
        trademarks.append({
            "商标名称": trademark_name_match.group(1).strip(),
            "类别": int(category_match.group(1)),
            "注册号": registration_number_match.group(1)
        })
   
    return {
        "文件名": filename, 
        "案件类型": case_type, 
        "申请人": applicant,
        "统一社会信用代码": unified_credit_code,
        "商标列表": trademarks
    }

def extract_opposition_case(text, filename):
    case_type = "商标异议"
    applicant = re.search(r'异议人名称：\s*([^\n]*?)\s+统一社会信用代码', 
                          text, re.IGNORECASE)
    applicant = applicant.group(1).strip() if applicant else "N/A"
    
    # 提取统一社会信用代码
    unified_credit_code_match = re.search(r'(?:统一社会信用代码|信用代码)[：:]\s*([0-9A-Z]{18})', text, re.IGNORECASE)
    unified_credit_code = unified_credit_code_match.group(1).strip() if unified_credit_code_match else "N/A"
    
    trademarks = []
    for m in re.finditer(r'被异议商标：\s*(.*?)\s+被异议类别：\s*(\d+).*?商标注册号：\s*([0-9A-Za-z]+)', 
                         text, re.DOTALL):
        trademarks.append({
            "商标名称": m.group(1).strip(), 
            "类别": int(m.group(2)), 
            "注册号": m.group(3)
        })
    
    return {
        "文件名": filename, 
        "案件类型": case_type, 
        "申请人": applicant,
        "统一社会信用代码": unified_credit_code,
        "商标列表": trademarks
    }

def extract_invalid_case(text, filename):
    case_type = "无效宣告"
    applicant = re.search(r'(?:申请人名称\$\$中文\$\$|申请人名称)：\s*([^\n]*?)(?=\s+(?:统一社会信用代码|地址))', 
                          text, re.DOTALL)
    applicant = applicant.group(1).strip() if applicant else "N/A"
    
    # 提取统一社会信用代码
    unified_credit_code_match = re.search(r'(?:统一社会信用代码|信用代码)[：:]\s*([0-9A-Z]{18})', text, re.IGNORECASE)
    unified_credit_code = unified_credit_code_match.group(1).strip() if unified_credit_code_match else "N/A"
    
    trademarks = []
    for m in re.finditer(r'争议商标：\s*(.*?)\s+类别：\s*(\d+).*?注册号/国际注册号：\s*([0-9A-Za-z]+)', 
                         text, re.DOTALL):
        trademarks.append({
            "商标名称": m.group(1).strip(), 
            "类别": int(m.group(2)), 
            "注册号": m.group(3)
        })
    
    return {
        "文件名": filename, 
        "案件类型": case_type, 
        "申请人": applicant,
        "统一社会信用代码": unified_credit_code,
        "商标列表": trademarks
    }

# ============================= 通用文档生成函数 =============================
def create_word_doc(applicant, records, output_dir, case_type):
    """生成Word请款单"""
    # 使用后台模板文件
    template_path = "请款单模板.docx"
    
    if not os.path.exists(template_path):
        st.error(f"错误: 找不到请款单模板文件 '{template_path}'")
        return None
    
    try:
        doc = Document(template_path)
        
        # 计算汇总
        if case_type == "新申请商标":
            case_types = ["商标注册申请"]
        else:
            case_types = list({r["案件类型"] for r in records})
        
        case_type_str = "、".join(case_types)
        total_official = sum(r["官费"] for r in records)
        total_agent = sum(r["代理费"] for r in records)
        total = total_official + total_agent
        
        # 替换正文占位符
        today_str = datetime.date.today().strftime("%Y年%m月%d日")
        for para in doc.paragraphs:
            for run in para.runs:
                run.text = run.text.replace("{申请人}", applicant) \
                                  .replace("{事宜类型}", case_type_str) \
                                  .replace("{日期}", today_str) \
                                  .replace("{总官费}", str(total_official)) \
                                  .replace("{总代理费}", str(total_agent)) \
                                  .replace("{总计}", str(total)) \
                                  .replace("{大写}", number_to_upper(total))
        
        # 动态写入表格
        if doc.tables:
            table = doc.tables[0]
            
            # 删除模板中的示例行（如果存在）
            if len(table.rows) > 1:
                for _ in range(len(table.rows) - 1, 0, -1):
                    table._tbl.remove(table.rows[1]._tr)
            
            # 添加数据行
            for idx, rec in enumerate(records, 1):
                row = table.add_row().cells
                row[0].text = str(idx)
                row[1].text = rec["案件类型"] if case_type != "新申请商标" else "商标注册申请"
                row[2].text = rec["商标名称"]
                row[3].text = str(rec["类别"])
                row[4].text = f"{rec['官费']}"
                row[5].text = f"{rec['代理费']}"
                row[6].text = f"{rec['官费'] + rec['代理费']}"
            
            # 追加合计行
            total_row = table.add_row().cells
            total_row[0].merge(total_row[3])
            total_row[0].text = "合计"
            total_row[4].text = f"{total_official}"
            total_row[5].text = f"{total_agent}"
            total_row[6].text = f"{total}"
        
        # 保存文件
        filename = f"请款单（{applicant}-{case_type_str}）-{total}-{datetime.date.today().strftime('%Y%m%d')}.docx"
        output_path = os.path.join(output_dir, filename)
        doc.save(output_path)
        
        return filename
    except Exception as e:
        st.error(f"生成Word文档时出错: {str(e)}")
        st.text(traceback.format_exc())
        return None

def build_excel(rows, output_dir):
    """生成Excel汇总表"""
    # 使用后台模板文件
    template_path = "发票申请表.xlsx"
    
    if not os.path.exists(template_path):
        st.error(f"错误: 找不到发票申请表模板文件 '{template_path}'")
        return None
    
    try:
        wb = load_workbook(template_path)
        ws = wb.active
        row_idx = 2
        
        for r in rows:
            ws[f"C{row_idx}"] = r["申请人"]
            ws[f"D{row_idx}"] = r["统一社会信用代码"]  # 统一社会信用代码列
            ws[f"G{row_idx}"] = r["总官费"]
            ws[f"H{row_idx}"] = r["总官费"]
            ws[f"I{row_idx}"] = r["总计"]
            ws[f"Q{row_idx}"] = datetime.date.today().strftime("%Y年%m月%d日")
            row_idx += 1
            
            ws[f"C{row_idx}"] = r["申请人"]
            ws[f"D{row_idx}"] = r["统一社会信用代码"]  # 统一社会信用代码列
            ws[f"G{row_idx}"] = r["总代理费"]
            ws[f"H{row_idx}"] = r["总代理费"]
            ws[f"I{row_idx}"] = r["总计"]
            ws[f"Q{row_idx}"] = datetime.date.today().strftime("%Y年%m月%d日")
            row_idx += 1
        
        excel_name = f"发票申请表-{datetime.date.today().strftime('%Y%m%d')}.xlsx"
        excel_path = os.path.join(output_dir, excel_name)
        wb.save(excel_path)
        
        return excel_name
    except Exception as e:
        st.error(f"生成Excel汇总时出错: {str(e)}")
        st.text(traceback.format_exc())
        return None

# ============================= 主应用逻辑 =============================
def main_app():
    # 案件类型选择
    st.header("1. 选择案件类型")
    st.session_state.case_type = st.radio(
        "请选择处理的案件类型:",
        ["新申请商标", "案件类商标"],
        index=0 if st.session_state.case_type == "新申请商标" else 1
    )
    
    case_type = st.session_state.case_type
    
    # 文件上传和处理区域
    st.header("2. 上传案件PDF文件")
    uploaded_files = st.file_uploader("请选择PDF文件", type="pdf", accept_multiple_files=True)

    if uploaded_files and st.button("处理PDF文件"):
        with st.spinner("正在处理PDF文件..."):
            try:
                # 创建临时目录
                temp_dir = tempfile.mkdtemp()
                st.session_state.temp_dir = temp_dir
                
                pdf_dir = os.path.join(temp_dir, "pdf_files")
                output_dir = os.path.join(temp_dir, "output")
                os.makedirs(pdf_dir, exist_ok=True)
                os.makedirs(output_dir, exist_ok=True)
                
                # 保存上传的文件
                for uploaded_file in uploaded_files:
                    file_path = os.path.join(pdf_dir, uploaded_file.name)
                    with open(file_path, "wb") as f:
                        f.write(uploaded_file.getbuffer())
                
                # 按申请人聚合
                applicant_map = defaultdict(list)
                extracted_data = []
                
                for filename in os.listdir(pdf_dir):
                    if filename.endswith(".pdf"):
                        try:
                            file_path = os.path.join(pdf_dir, filename)
                            
                            if case_type == "新申请商标":
                                # 新申请商标处理
                                data = extract_pdf_data(file_path)
                                applicant = data["申请人"]
                                unified_credit_code = data["统一社会信用代码"]
                                
                                for tm in data["商标列表"]:
                                    # 处理需要手动输入的类别
                                    if tm["类别"] == "MANUAL_INPUT_REQUIRED":
                                        # 在后续步骤中处理
                                        continue
                                    
                                    applicant_map[applicant].append({
                                        "商标名称": tm["商标名称"],
                                        "类别": tm["类别"],
                                        "案件类型": "商标注册申请",
                                        "官费": OFFICIAL_FEES["新申请商标"],
                                        "统一社会信用代码": unified_credit_code,  # 添加统一社会信用代码
                                    })
                                
                                extracted_data.append(data)
                                st.success(f"成功处理: {filename} (申请人: {applicant})")
                            
                            else:
                                # 案件类商标处理
                                with pdfplumber.open(file_path) as pdf:
                                    text = []
                                    for page in pdf.pages:
                                        txt = page.extract_text()
                                        if not txt:
                                            continue
                                        if any(k in txt for k in ["申请书", "申 请 书", "撤销", "异议", "无效", "宣告"]):
                                            txt = txt.replace("　", " ").replace("\xa0", " ")
                                            txt = re.sub(r'[\u3000]', ' ', txt)
                                            text.append(txt)
                                    text = "".join(text).strip()
                                
                                # 提取案件信息
                                data = extract_case_info(text, filename)
                                applicant = data["申请人"]
                                unified_credit_code = data["统一社会信用代码"]
                                
                                for tm in data["商标列表"]:
                                    applicant_map[applicant].append({
                                        "商标名称": tm["商标名称"],
                                        "类别": tm["类别"],
                                        "案件类型": data["案件类型"],
                                        "官费": OFFICIAL_FEES[data["案件类型"]],
                                        "统一社会信用代码": unified_credit_code,  # 添加统一社会信用代码
                                    })
                                
                                extracted_data.append(data)
                                st.success(f"成功处理: {filename} (申请人: {applicant}, 类型: {data['案件类型']})")
                                
                        except Exception as e:
                            st.error(f"处理文件 {filename} 时出错: {str(e)}")
                            st.text(traceback.format_exc())
                
                # 保存处理结果到session
                st.session_state.extracted_data = extracted_data
                st.session_state.applicant_map = dict(applicant_map)
                st.session_state.processing_stage = 1
                
                st.success(f"成功处理 {len(uploaded_files)} 个PDF文件！")
                st.info(f"共发现 {len(applicant_map)} 个申请人")
                
            except Exception as e:
                st.error(f"处理过程中发生错误: {str(e)}")
                st.text(traceback.format_exc())

    # 显示提取结果
    if st.session_state.processing_stage >= 1 and st.session_state.extracted_data:
        st.header("3. 提取结果")
        
        for applicant, records in st.session_state.applicant_map.items():
            with st.expander(f"申请人: {applicant}"):
                st.write(f"统一社会信用代码: {records[0].get('统一社会信用代码', 'N/A')}")
                st.write(f"案件数量: {len(records)}")
                for record in records:
                    st.write(f"- 商标: {record['商标名称']}, 类别: {record['类别']}, 类型: {record['案件类型']}, 官费: {record['官费']}元")
                
                # 显示新申请商标需要手动输入的类别
                if case_type == "新申请商标":
                    for data in st.session_state.extracted_data:
                        if data["申请人"] == applicant:
                            for tm in data["商标列表"]:
                                if tm["类别"] == "MANUAL_INPUT_REQUIRED":
                                    st.warning(f"商标 '{tm['商标名称']}' 需要手动输入类别")

    # 设置代理费和手动输入类别
    if st.session_state.processing_stage >= 1 and st.session_state.applicant_map:
        st.header("4. 设置参数")
        
        # 设置代理费
        st.subheader("代理费设置")
        for applicant in st.session_state.applicant_map.keys():
            default_fee = st.session_state.agent_fees.get(applicant, 600)
            fee = st.number_input(
                f"{applicant}的代理费(元/件)", 
                min_value=0, 
                value=default_fee,
                key=f"fee_{applicant}"
            )
            st.session_state.agent_fees[applicant] = fee
        
        # 新申请商标需要手动输入类别
        if case_type == "新申请商标":
            st.subheader("商标类别设置")
            for data in st.session_state.extracted_data:
                applicant = data["申请人"]
                for tm in data["商标列表"]:
                    if tm["类别"] == "MANUAL_INPUT_REQUIRED":
                        key = f"manual_{applicant}_{tm['商标名称']}"
                        categories = st.text_input(
                            f"商标 '{tm['商标名称']}' 的类别(多个类别用逗号分隔)", 
                            key=key,
                            placeholder="例如: 9,35,42"
                        )
                        
                        # 保存手动输入的类别
                        if categories:
                            st.session_state[key] = categories

    # 生成文档按钮
    if st.session_state.processing_stage >= 1 and st.session_state.applicant_map and st.button("生成请款单"):
        with st.spinner("正在生成请款单和汇总表..."):
            try:
                output_dir = os.path.join(st.session_state.temp_dir, "output")
                os.makedirs(output_dir, exist_ok=True)
                
                generated_files = []
                excel_rows = []
                
                for applicant, records in st.session_state.applicant_map.items():
                    try:
                        # 添加代理费到记录
                        agent_fee = st.session_state.agent_fees.get(applicant, 1000)
                        
                        # 对于新申请商标，添加手动输入的类别
                        if st.session_state.case_type == "新申请商标":
                            # 创建一个新的记录列表，处理手动输入
                            processed_records = []
                            unified_credit_code = records[0].get("统一社会信用代码", "N/A")
                            
                            for data in st.session_state.extracted_data:
                                if data["申请人"] == applicant:
                                    for tm in data["商标列表"]:
                                        if tm["类别"] == "MANUAL_INPUT_REQUIRED":
                                            key = f"manual_{applicant}_{tm['商标名称']}"
                                            categories_input = st.session_state.get(key, "")
                                            if categories_input:
                                                categories = [cat.strip() for cat in categories_input.split(",") if cat.strip()]
                                                for cat in categories:
                                                    processed_records.append({
                                                        "商标名称": tm["商标名称"],
                                                        "类别": cat,
                                                        "案件类型": "商标注册申请",
                                                        "官费": OFFICIAL_FEES["新申请商标"],
                                                        "代理费": agent_fee,
                                                        "统一社会信用代码": unified_credit_code,
                                                    })
                                        else:
                                            processed_records.append({
                                                "商标名称": tm["商标名称"],
                                                "类别": tm["类别"],
                                                "案件类型": "商标注册申请",
                                                "官费": OFFICIAL_FEES["新申请商标"],
                                                "代理费": agent_fee,
                                                "统一社会信用代码": unified_credit_code,
                                            })
                        else:
                            # 案件类商标直接添加代理费
                            processed_records = []
                            unified_credit_code = records[0].get("统一社会信用代码", "N/A")
                            
                            for record in records:
                                record["代理费"] = agent_fee
                                record["统一社会信用代码"] = unified_credit_code
                                processed_records.append(record)
                        
                        # 生成Word文档
                        if processed_records:
                            word_filename = create_word_doc(
                                applicant, 
                                processed_records, 
                                output_dir,
                                st.session_state.case_type
                            )
                            
                            if word_filename:
                                word_path = os.path.join(output_dir, word_filename)
                                with open(word_path, "rb") as f:
                                    word_data = f.read()
                                
                                generated_files.append({
                                    "name": word_filename,
                                    "data": word_data,
                                    "type": "word"
                                })
                                
                                # 收集汇总数据
                                total_official = sum(r["官费"] for r in processed_records)
                                total_agent = sum(r["代理费"] for r in processed_records)
                                excel_rows.append({
                                    "申请人": applicant,
                                    "统一社会信用代码": unified_credit_code,  # 添加统一社会信用代码
                                    "总官费": total_official,
                                    "总代理费": total_agent,
                                    "总计": total_official + total_agent,
                                })
                    
                    except Exception as e:
                        st.error(f"为申请人 '{applicant}' 生成请款单时出错: {str(e)}")
                        st.text(traceback.format_exc())
                
                # 生成Excel汇总
                if excel_rows:
                    excel_filename = build_excel(excel_rows, output_dir)
                    if excel_filename:
                        excel_path = os.path.join(output_dir, excel_filename)
                        with open(excel_path, "rb") as f:
                            excel_data = f.read()
                        
                        generated_files.append({
                            "name": excel_filename,
                            "data": excel_data,
                            "type": "excel"
                        })
                
                # 保存生成的文件到session
                st.session_state.generated_files = generated_files
                st.session_state.processing_stage = 2
                st.success("文档生成完成！")
            except Exception as e:
                st.error(f"生成过程中发生错误: {str(e)}")
                st.text(traceback.format_exc())

    # 下载区域
    if st.session_state.processing_stage == 2 and st.session_state.generated_files:
        st.header("5. 下载生成的文件")
        
        # 显示所有生成的文件
        st.subheader("生成的文件列表")
        
        word_files = [f for f in st.session_state.generated_files if f["type"] == "word"]
        excel_files = [f for f in st.session_state.generated_files if f["type"] == "excel"]
        
        if word_files:
            st.subheader("请款单")
            for file in word_files:
                st.download_button(
                    label=f"下载 {file['name']}",
                    data=file["data"],
                    file_name=file["name"],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        if excel_files:
            st.subheader("汇总表")
            for file in excel_files:
                st.download_button(
                    label=f"下载 {file['name']}",
                    data=file["data"],
                    file_name=file["name"],
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

    # 重置按钮
    if st.button("重置所有数据"):
        # 清除所有session状态
        keys_to_clear = list(st.session_state.keys())
        for key in keys_to_clear:
            if key != 'temp_dir' and key != 'case_type':  # 保留temp_dir和case_type
                del st.session_state[key]
        
        # 清理临时目录
        if st.session_state.temp_dir and os.path.exists(st.session_state.temp_dir):
            try:
                shutil.rmtree(st.session_state.temp_dir)
            except:
                pass
        
        # 重新初始化必要的状态
        st.session_state.processing_stage = 0
        st.session_state.extracted_data = None
        st.session_state.agent_fees = {}
        st.session_state.generated_files = []
        st.session_state.temp_dir = ""
        
        st.success("系统已重置，可以开始新的处理流程！")

# ============================= 应用入口 =============================
# 显示模板状态
st.sidebar.header("系统状态")
payment_template_exists = os.path.exists("请款单模板.docx")
invoice_template_exists = os.path.exists("发票申请表.xlsx")

if payment_template_exists and invoice_template_exists:
    st.sidebar.success("✅ 模板文件已就绪")
    st.sidebar.info("请款单模板: 请款单模板.docx")
    st.sidebar.info("发票申请表模板: 发票申请表.xlsx")
    main_app()
else:
    st.sidebar.error("⚠️ 模板文件缺失")
    if not payment_template_exists:
        st.sidebar.error("请款单模板 '请款单模板.docx' 不存在")
    if not invoice_template_exists:
        st.sidebar.error("发票申请表模板 '发票申请表.xlsx' 不存在")
    
    st.error("系统无法启动，因为缺少必要的模板文件。请确保以下文件与应用程序在同一目录下:")
    st.error("- 请款单模板.docx")
    st.error("- 发票申请表.xlsx")
    
    st.info("请上传模板文件后重新启动应用程序")
